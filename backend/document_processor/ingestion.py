"""
Module for document ingestion and processing.
Handles loading, splitting and storing documents from various sources.
"""

from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_community.document_loaders import (
    WebBaseLoader,
    PyPDFLoader,
    TextLoader,
    DirectoryLoader
)
from langchain_openai import OpenAIEmbeddings
from typing import List, Any, Union
import os
from langchain_core.documents import Document
from .interfaces import DocumentLoader, TextSplitter, VectorStore
import docx
import psutil
import shutil

load_dotenv()

class WebLoader(DocumentLoader):
    """Loads documents from web URLs."""
    
    def __init__(self, urls: List[str]):
        """
        Initialize web loader.
        
        Args:
            urls: List of URLs to load
        """
        self.urls = urls

    def load(self) -> List[Document]:
        """Load documents from URLs."""
        docs = [WebBaseLoader(url).load() for url in self.urls]
        return [item for sublist in docs for item in sublist]

class PDFLoader(DocumentLoader):
    """Loads documents from PDF files."""
    
    def __init__(self, pdf_files: List[str]):
        """
        Initialize PDF loader.
        
        Args:
            pdf_files: List of PDF file paths
        """
        self.pdf_files = pdf_files

    def load(self) -> List[Document]:
        """Load documents from PDF files."""
        docs = []
        for pdf in self.pdf_files:
            if os.path.exists(pdf):
                loader = PyPDFLoader(pdf)
                docs.extend(loader.load())
        return docs

class FileLoader(DocumentLoader):
    """Loads documents from text files."""
    
    def __init__(self, text_files: List[str]):
        """
        Initialize text file loader.
        
        Args:
            text_files: List of text file paths
        """
        self.text_files = text_files

    def load(self) -> List[Document]:
        """Load documents from text files."""
        docs = []
        for text_file in self.text_files:
            if os.path.exists(text_file):
                loader = TextLoader(text_file)
                docs.extend(loader.load())
        return docs

class DirectoryDocumentLoader(DocumentLoader):
    """Loads all documents from a directory."""
    
    def __init__(self, directory_path: str, glob_pattern: str = "**/*"):
        """
        Initialize directory loader.
        
        Args:
            directory_path: Path to directory
            glob_pattern: Pattern for file matching
        """
        self.directory_path = directory_path
        self.glob_pattern = glob_pattern

    def load(self) -> List[Document]:
        """Load all documents from directory."""
        loader = DirectoryLoader(
            self.directory_path,
            glob=self.glob_pattern,
            use_multithreading=True
        )
        return loader.load()

class DocxLoader(DocumentLoader):
    """Loads documents from DOCX files."""
    
    def __init__(self, docx_files: List[str]):
        """
        Initialize DOCX loader.
        
        Args:
            docx_files: List of DOCX file paths
        """
        self.docx_files = docx_files

    def load(self) -> List[Document]:
        """Load documents from DOCX files."""
        docs = []
        for file_path in self.docx_files:
            if os.path.exists(file_path):
                doc = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path}
                ))
        return docs

class RecursiveTextSplitter(TextSplitter):
    """Splits documents into smaller chunks recursively."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """
        Initialize text splitter.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ";", ":", ",", " ", ""],
            length_function=len
        )

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks."""
        return self.splitter.split_documents(documents)

class ChromaVectorStore(VectorStore):
    def __init__(self, collection_name: str = "rag-chroma", persist_directory: str = "./.chroma"):
        self.collection_name = collection_name
        self.persist_directory = persist_directory
        self.embedding_function = OpenAIEmbeddings()
        self.vectorstore = None
        self._client = None

    def _get_client(self):
        """Create a new client instance"""
        from chromadb import Client
        from chromadb.config import Settings
        
        if not self._client:
            self._client = Client(Settings(
                persist_directory=self.persist_directory,
                is_persistent=True
            ))
        return self._client

    def store_documents(self, documents: List[Document]) -> Any:
        client = self._get_client()

        # Create and return the vectorstore
        self.vectorstore = Chroma.from_documents(
            documents=documents,
            collection_name=self.collection_name,
            embedding=self.embedding_function,
            persist_directory=self.persist_directory,
            client=client
        )
        return self.vectorstore

    def get_retriever(self) -> Any:
        if not self.vectorstore:
            client = self._get_client()

            self.vectorstore = Chroma(
                collection_name=self.collection_name,
                persist_directory=self.persist_directory,
                embedding_function=self.embedding_function,
                client=client
            )
        return self.vectorstore.as_retriever()

    def cleanup(self):
        """Clean up vector store directory"""
        try:
            # First, try to delete the collection
            if self.vectorstore is not None and hasattr(self.vectorstore, '_collection'):
                # Delete all documents using a valid where clause
                self.vectorstore._collection.delete(where={"source": {"$ne": ""}})

            # Then wait a bit
            import time
            time.sleep(1)

            # Finally remove the directory
            if os.path.exists(self.persist_directory):
                # Get current process
                current_process = psutil.Process()
                
                # Close all file handles in the directory
                for handler in current_process.open_files():
                    if self.persist_directory in handler.path:
                        try:
                            os.close(handler.fd)
                        except:
                            pass

                # Remove directory
                shutil.rmtree(self.persist_directory, ignore_errors=True)
                print(f"Successfully cleaned up vector store at {self.persist_directory}")
        except Exception as e:
            print(f"Error cleaning up vector store: {str(e)}")

class DocumentIngester:
    """Handles document ingestion workflow."""
    
    def __init__(
        self,
        text_splitter: TextSplitter,
        vector_store: VectorStore
    ):
        """
        Initialize document ingester.
        
        Args:
            text_splitter: Splitter for chunking documents
            vector_store: Store for document vectors
        """
        self.text_splitter = text_splitter
        self.vector_store = vector_store

    def process_documents(self, document_loader: DocumentLoader) -> Any:
        """
        Process and store documents.
        
        Args:
            document_loader: Loader for documents
            
        Returns:
            Stored documents in vector store
        """
        documents = document_loader.load()
        split_docs = self.text_splitter.split_documents(documents)
        return self.vector_store.store_documents(split_docs)

class CombinedLoader(DocumentLoader):
    """Combines multiple document loaders into one."""
    
    def __init__(self, loaders: List[DocumentLoader]):
        """
        Initialize combined loader.
        
        Args:
            loaders: List of document loaders to combine
        """
        self.loaders = loaders

    def load(self) -> List[Document]:
        """Load documents from all loaders."""
        documents = []
        for loader in self.loaders:
            documents.extend(loader.load())
        return documents

def get_document_loader(file_paths: List[str]) -> DocumentLoader:
    """
    Get appropriate loader(s) for file types.
    
    Args:
        file_paths: List of file paths
        
    Returns:
        Document loader that can handle all provided files
    """
    if not file_paths:
        raise ValueError("No files provided")

    # Group files by extension
    pdf_files = [f for f in file_paths if f.lower().endswith('.pdf')]
    docx_files = [f for f in file_paths if f.lower().endswith('.docx')]
    txt_files = [f for f in file_paths if f.lower().endswith('.txt')]
    
    # Create loaders for each file type
    loaders = []
    if pdf_files:
        loaders.append(PDFLoader(pdf_files))
    if docx_files:
        loaders.append(DocxLoader(docx_files))
    if txt_files:
        loaders.append(FileLoader(txt_files))
        
    if not loaders:
        raise ValueError(f"Unsupported file type(s). Supported types are: .pdf, .docx, .txt")
        
    # If only one loader, return it directly
    if len(loaders) == 1:
        return loaders[0]
        
    # Otherwise, combine all loaders
    return CombinedLoader(loaders)